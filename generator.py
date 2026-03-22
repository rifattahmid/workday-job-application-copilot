import os
import re
import shutil
import anthropic
from datetime import datetime
from docx import Document
from docx2pdf import convert
from dotenv import load_dotenv

from config import OUTPUT_BASE, TEMPLATE_BASE

load_dotenv()

DATE_PATTERN = re.compile(r"\d{1,2}\s+\w+\s+20\d{2}")


def _set_para_text(para, text):
    """Replace paragraph text preserving the first run's formatting (used for dates)."""
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def classify_job(title, description):
    """Match job to a template subfolder. Reads available folders dynamically.

    Scoring:
    - Title keyword match  → +3  (title signals intent; weighted 3x over body text)
    - Description keyword  → +1
    Tie-break order: finance > investment > accounting.
    """
    title_text = title.lower()
    desc_text  = description.lower()
    available  = [
        f for f in os.listdir(TEMPLATE_BASE)
        if os.path.isdir(os.path.join(TEMPLATE_BASE, f))
    ]

    # Keyword map — extend as new folders are added
    keywords = {
        "investment": ["investment", "investor", "portfolio", "asset management",
                       "acquisition", "acquisitions", "mergers", "m&a",
                       "transaction", "deals", "fund", "equity", "credit",
                       "real estate", "real assets", "infrastructure", "capital markets"],
        "accounting": ["account", "audit", "tax", "bookkeep", "controller",
                       "cpa", "chartered accountant"],
        "finance":    ["finance", "financial", "fp&a", "treasury", "budget",
                       "forecast", "analyst", "corporate finance", "m&a",
                       "mergers", "acquisitions", "transaction", "deals", "valuation"],
    }

    TITLE_MULTIPLIER = 3

    scores = {folder: 0 for folder in available}
    for folder in available:
        for kw in keywords.get(folder, []):
            if kw in title_text:
                scores[folder] += TITLE_MULTIPLIER
            if kw in desc_text:
                scores[folder] += 1

    best = max(scores, key=lambda f: scores[f]) if scores else available[0]

    # Tie-break: if two or more categories share the top score, prefer finance > investment > accounting
    top_score = scores[best]
    if sum(1 for v in scores.values() if v == top_score) > 1:
        for preferred in ("finance", "investment", "accounting"):
            if preferred in available and scores.get(preferred, -1) == top_score:
                best = preferred
                break

    print(f"  Job classified as: {best}  (scores: {scores})")
    return best


def get_paths(category):
    base = os.path.join(TEMPLATE_BASE, category)
    if not os.path.isdir(base):
        raise FileNotFoundError(f"Template folder not found: {base}")

    resume_pdf = cover_docx = resume_txt = None
    for f in os.listdir(base):
        f_lower = f.lower()
        if f_lower.endswith(".pdf") and "resume" in f_lower:
            resume_pdf = os.path.join(base, f)
        elif f_lower.endswith(".docx") and "cover" in f_lower:
            cover_docx = os.path.join(base, f)
        elif f_lower.endswith(".txt") and "resume" in f_lower:
            resume_txt = os.path.join(base, f)

    missing = [
        name for name, val in [
            ("Resume.pdf", resume_pdf),
            ("Cover Letter.docx", cover_docx),
            ("Resume.txt", resume_txt),
        ]
        if val is None
    ]
    if missing:
        raise FileNotFoundError(f"Missing in {base}: {', '.join(missing)}")
    return resume_pdf, cover_docx, resume_txt



def _rebold_title(para, title):
    """After _set_para_text, split run[0] around `title` and make the title portion bold."""
    if not para.runs:
        return
    text = para.runs[0].text
    if title not in text:
        return
    idx = text.index(title)
    before = text[:idx]
    after = text[idx + len(title):]
    run0 = para.runs[0]
    # Copy font properties for new runs
    font_name = run0.font.name
    font_size = run0.font.size
    run0.text = before
    bold_run = para.add_run(title)
    bold_run.bold = True
    bold_run.font.name = font_name
    bold_run.font.size = font_size
    if after:
        after_run = para.add_run(after)
        after_run.bold = False
        after_run.font.name = font_name
        after_run.font.size = font_size


def fill_cover_letter(path, company, title, intro, responsibilities, qualifications):
    doc = Document(path)
    today = datetime.now().strftime("%d %B %Y")

    # Replace dates (checks full paragraph text to catch cross-run dates)
    for para in doc.paragraphs:
        if DATE_PATTERN.search(para.text):
            _set_para_text(para, DATE_PATTERN.sub(today, para.text))

    # Collect paragraphs with _ blanks and note if any run in that para was bold
    blank_paras = [
        (i, para, any(r.bold and "_" in r.text for r in para.runs))
        for i, para in enumerate(doc.paragraphs) if "_" in para.text
    ]

    print("\n=== BLANKS FOUND IN TEMPLATE ===")
    for _, para, _ in blank_paras:
        print(f"  {repr(para.text)}")
    print("=================================\n")

    if not blank_paras:
        doc.save(path)
        print(f"  Cover letter saved: {path}")
        return

    numbered_lines = "\n".join(
        f"{j+1}. {para.text}" for j, (_, para, _) in enumerate(blank_paras)
    )

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=1000,
        messages=[{"role": "user", "content": f"""Rewrite each cover letter sentence below, filling every blank (_) and adjusting surrounding words for natural flow.

Role: {title}
Company (USE EXACTLY THIS): {company}
Today: {today}

Job description: {intro}
Responsibilities: {responsibilities}
Qualifications: {qualifications}

Sentences to rewrite:
{numbered_lines}

Rules:
- Use EXACTLY "{company}" for the company name — never modify or expand it
- Use EXACTLY "{title}" for the role/position — never modify it
- For region: be specific to this job (e.g. "Australian market" if the role is in Australia)
- You have FULL CREATIVE FREEDOM to rewrite the ENTIRE sentence so it flows naturally and is specific to this role and asset class
- Match terminology to the asset class: real estate/property → "assets", "properties", "real estate portfolio"; credit → "credit investments", "debt instruments"; infrastructure → "assets", "projects"; private equity → "portfolio companies"; equities → "companies", "equity portfolio"
- NEVER use private equity language (e.g. "portfolio companies", "value creation alongside companies") for real estate, infrastructure, or credit roles
- If the sentence ends awkwardly after filling the blank, rewrite the ending completely
- Return ONLY the rewritten sentences, numbered the same way"""}]
    )

    response = message.content[0].text.strip()
    print(f"  Claude rewrites:\n{response}\n")

    rewrites = {}
    for line in response.splitlines():
        m = re.match(r'^(\d+)\.\s*(.+)$', line.strip())
        if m:
            rewrites[int(m.group(1)) - 1] = m.group(2).strip()

    for j, (_, para, had_bold_blank) in enumerate(blank_paras):
        new_text = rewrites.get(j, "")
        if not new_text:
            continue
        _set_para_text(para, new_text)
        if title in new_text and para.runs:
            _rebold_title(para, title)

    doc.save(path)
    print(f"  Cover letter saved: {path}")


def generate_application(data):
    title = data["title"]
    company = data["company"]

    category = classify_job(title, data["description"])
    resume_pdf, cover_docx, resume_txt = get_paths(category)

    folder_name = f"{company} - {title}".replace("/", "-")
    output_folder = os.path.join(OUTPUT_BASE, folder_name)
    os.makedirs(output_folder, exist_ok=True)

    shutil.copy(resume_pdf, os.path.join(output_folder, os.path.basename(resume_pdf)))
    cover_dest = os.path.join(output_folder, os.path.basename(cover_docx))
    shutil.copy(cover_docx, cover_dest)

    # Write position description PDF (captured during scraping while session was live)
    position_pdf = os.path.join(output_folder, "Position Description.pdf")
    if data.get("pdf_bytes"):
        with open(position_pdf, "wb") as f:
            f.write(data["pdf_bytes"])
        print(f"  Position description PDF: {position_pdf}")

    with open(resume_txt, "r", encoding="utf-8") as f:
        resume_text = f.read()  # noqa: F841 (available for future use)

    fill_cover_letter(
        cover_dest, company, title,
        data.get("intro", ""),
        data.get("responsibilities", ""),
        data.get("qualifications", ""),
    )

    pdf_dest = cover_dest.replace(".docx", ".pdf")
    try:
        convert(cover_dest, pdf_dest)
        print(f"  Cover letter PDF: {pdf_dest}")
    except Exception as e:
        print(f"  WARNING: PDF conversion failed ({e})")
        print("  Make sure Microsoft Word is installed and the .docx is not open.")

    print(f"\nDone! Saved to: {output_folder}")
    os.startfile(output_folder)
    return output_folder
